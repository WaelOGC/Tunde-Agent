"""
Generate professional exports from saved HTML reports (PDF, Word, CSV, Tailwind HTML).

Used by Telegram post-report actions and optional email attachments.
"""

from __future__ import annotations

import base64
import csv
import html
import io
import json
import logging
import re
from io import BytesIO
from typing import Any

from tunde_agent.services.report_html import reports_dir

logger = logging.getLogger(__name__)


def read_report_html(report_id: str) -> str | None:
    rid = report_id.strip()
    path = reports_dir() / f"{rid}.html"
    if not path.is_file():
        return None
    return path.read_text(encoding="utf-8")


def html_to_plain(raw_html: str, *, max_chars: int = 80_000) -> str:
    t = re.sub(r"(?is)<script.*?>.*?</script>", " ", raw_html)
    t = re.sub(r"(?is)<style.*?>.*?</style>", " ", t)
    t = re.sub(r"(?is)<[^>]+>", " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t[:max_chars]


def extract_title_from_html(raw_html: str, *, fallback: str = "Tunde AI Agent report") -> str:
    m = re.search(r"<title[^>]*>([^<]+)</title>", raw_html, re.I | re.DOTALL)
    if not m:
        return fallback
    title = re.sub(r"\s+", " ", m.group(1)).strip()
    title = re.sub(r"\s*·\s*Tunde.*$", "", title, flags=re.I).strip()
    return title[:500] or fallback


def _pdf_safe(text: str) -> str:
    return text.encode("latin-1", errors="replace").decode("latin-1")


def build_pdf_bytes(title: str, body_plain: str) -> bytes:
    """Legacy simple PDF (tests); prefer ``build_professional_pdf_from_report`` for deliveries."""
    from fpdf import FPDF

    pdf = FPDF(unit="mm", format="A4")
    pdf.set_margins(18, 18, 18)
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    epw = pdf.epw
    pdf.set_font("Helvetica", "B", 14)
    pdf.multi_cell(epw, 8, _pdf_safe(title[:500]))
    pdf.ln(4)
    pdf.set_font("Helvetica", "", 10)
    for para in body_plain.split("\n\n"):
        p = _pdf_safe(para.strip())
        if not p:
            continue
        pdf.multi_cell(epw, 5.5, p)
        pdf.ln(2)
    out = pdf.output()
    if isinstance(out, str):
        return out.encode("latin-1")
    return bytes(out)


def load_report_meta(report_id: str) -> dict[str, Any]:
    path = reports_dir() / f"{report_id.strip()}.meta.json"
    if not path.is_file():
        return {}
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
        return data if isinstance(data, dict) else {}
    except Exception:
        logger.warning("Could not read report meta for %s", report_id[:16])
        return {}


def _resolve_chart_png_for_pdf(meta: dict[str, Any]) -> bytes | None:
    b64 = meta.get("chart_png_b64")
    if isinstance(b64, str) and b64.strip():
        try:
            raw = base64.standard_b64decode(b64.strip())
            if len(raw) > 80:
                return raw
        except Exception:
            pass
    cm = meta.get("analyst_chart_metrics")
    if isinstance(cm, dict):
        from tunde_agent.services.research_orchestration.designer_agent import (
            generate_charts_from_analyst_with_fallback,
        )

        topic = str(meta.get("topic") or "report")
        charts = generate_charts_from_analyst_with_fallback(
            {"chart_metrics": cm},
            theme_topic=topic,
        )
        if charts:
            return charts[0][0]
    return None


def build_professional_pdf_from_report(report_id: str) -> bytes:
    """
    A4 PDF with safe margins, embedded chart image when available, narrative, and clickable resource links.
    """
    from fpdf import FPDF

    meta = load_report_meta(report_id)
    loaded = load_report_for_export(report_id)
    title = "Tunde AI Agent report"
    plain = ""
    if loaded:
        title, plain = loaded
    if meta.get("topic"):
        title = str(meta["topic"])[:300]

    pdf = FPDF(unit="mm", format="A4")
    pdf.set_margins(18, 18, 18)
    pdf.set_auto_page_break(auto=True, margin=22)
    pdf.add_page()
    epw = pdf.epw

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(epw, 8, _pdf_safe(title))
    pdf.ln(2)
    pdf.set_font("Helvetica", "I", 9)
    pdf.set_text_color(80, 80, 80)
    pdf.multi_cell(
        epw,
        4.5,
        _pdf_safe("Tunde AI Agent — figures include labeled estimates where sources are thin."),
    )
    pdf.set_text_color(0, 0, 0)
    pdf.ln(3)

    chart_png = _resolve_chart_png_for_pdf(meta)
    if chart_png:
        try:
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 7, _pdf_safe("Visual analytics"))
            pdf.ln(6)
            pdf.image(io.BytesIO(chart_png), w=min(epw, 175), type="PNG")
            pdf.ln(4)
        except Exception as exc:
            logger.warning("PDF chart embed skipped: %s", exc)

    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 7, _pdf_safe("Report narrative"))
    pdf.ln(5)
    pdf.set_font("Helvetica", "", 10)
    for para in plain.split("\n\n"):
        p = _pdf_safe(para.strip())
        if not p:
            continue
        pdf.multi_cell(epw, 5.2, p)
        pdf.ln(1.2)

    pdf.ln(2)
    pdf.set_font("Helvetica", "B", 12)
    pdf.cell(0, 7, _pdf_safe("Resources"))
    pdf.ln(5)
    pdf.set_font("Helvetica", "", 9)

    ru = (meta.get("report_url") or "").strip()
    if ru.startswith("http"):
        pdf.set_text_color(0, 60, 170)
        pdf.cell(0, 5, _pdf_safe("Open interactive report"), link=ru[:2048])
        pdf.ln(6)
        pdf.set_text_color(0, 0, 0)

    gs = "https://docs.google.com/spreadsheets/u/0/create"
    pdf.set_text_color(0, 60, 170)
    pdf.cell(0, 5, _pdf_safe("Google Sheets — create / import data"), link=gs)
    pdf.ln(6)
    pdf.set_text_color(0, 0, 0)

    for s in meta.get("sources") or []:
        if not isinstance(s, dict):
            continue
        u = (s.get("url") or "").strip()
        t = (s.get("title") or "Source").strip()[:100]
        if u.startswith("http"):
            pdf.set_text_color(0, 60, 170)
            pdf.cell(0, 5, _pdf_safe(t), link=u[:2048])
            pdf.ln(5)
            pdf.set_text_color(0, 0, 0)

    out = pdf.output()
    if isinstance(out, str):
        return out.encode("latin-1")
    return bytes(out)


def build_docx_bytes(title: str, body_plain: str) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(title[:200], level=0)
    for para in body_plain.split("\n\n"):
        t = para.strip()
        if t:
            doc.add_paragraph(t[:8000])
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_sheet_csv_bytes(title: str, body_plain: str) -> bytes:
    """UTF-8 with BOM for Excel; import into Google Sheets via File → Import."""
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["report_title", title])
    w.writerow(["block", "content"])
    step = 3500
    for i in range(0, len(body_plain), step):
        chunk = body_plain[i : i + step].replace("\r", " ")
        w.writerow([f"block_{i // step + 1}", chunk])
    return buf.getvalue().encode("utf-8-sig")


def build_tailwind_landing_html(
    title: str,
    body_plain: str,
    *,
    report_url: str = "",
) -> str:
    """Standalone page using Tailwind CDN (clean, shareable)."""
    safe_title = html.escape(title[:300])
    paras = [p.strip() for p in body_plain.split("\n\n") if p.strip()][:120]
    blocks = "".join(
        f'<p class="mb-5 text-slate-600 dark:text-slate-300 leading-relaxed text-lg">{html.escape(p[:6000])}</p>'
        for p in paras
    )
    link_row = ""
    if (report_url or "").strip():
        u = html.escape(report_url.strip()[:2048], quote=True)
        link_row = (
            f'<p class="mt-8"><a href="{u}" class="text-indigo-600 dark:text-indigo-400 font-medium underline">'
            "Open full interactive report</a></p>"
        )
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8"/>
  <meta name="viewport" content="width=device-width, initial-scale=1"/>
  <title>{safe_title}</title>
  <script src="https://cdn.tailwindcss.com"></script>
</head>
<body class="min-h-screen bg-slate-50 dark:bg-slate-950 text-slate-900 dark:text-slate-100">
  <main class="max-w-3xl mx-auto px-6 py-14">
    <p class="text-xs uppercase tracking-widest text-indigo-500 dark:text-indigo-400 mb-3">Tunde AI Agent</p>
    <h1 class="text-4xl font-bold tracking-tight mb-10 text-slate-900 dark:text-white">{safe_title}</h1>
    <article class="prose prose-lg dark:prose-invert max-w-none">
      {blocks}
    </article>
    {link_row}
    <footer class="mt-16 pt-8 border-t border-slate-200 dark:border-slate-800 text-sm text-slate-500">
      Copyright © 2026 Tunde AI Agent · NewFinity
    </footer>
  </main>
</body>
</html>
"""


def load_report_for_export(report_id: str) -> tuple[str, str] | None:
    """Return ``(title, plain_body)`` or ``None`` if missing."""
    raw = read_report_html(report_id)
    if raw is None:
        return None
    title = extract_title_from_html(raw)
    plain = html_to_plain(raw)
    if len(plain) < 20:
        return None
    return title, plain
